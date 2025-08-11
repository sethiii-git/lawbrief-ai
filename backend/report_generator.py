# File: backend/report_generator.py
"""
Professional PDF and DOCX report generation module.
"""

import io
import logging
import os
from datetime import datetime
from typing import Dict, List
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-interactive backend

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
from reportlab.platypus.tableofcontents import TableOfContents
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.shared import OxmlElement, qn


def generate_pdf_report(
    output_path: str,
    metadata: Dict,
    clauses: List[Dict], 
    risk_summary: Dict,
    summaries: Dict
) -> None:
    """
    Generate a professional PDF report of contract analysis.
    
    Args:
        output_path: Path to save the PDF report
        metadata: Contract metadata (filename, upload_date, etc.)
        clauses: List of extracted clauses
        risk_summary: Risk assessment results
        summaries: Document summaries
    """
    try:
        # Create document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=18
        )
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            spaceBefore=20
        )
        
        # Build content
        story = []
        
        # Title page
        story.append(Paragraph("LawBrief AI - Contract Analysis Report", title_style))
        story.append(Spacer(1, 0.3*inch))
        
        # Metadata section
        story.append(Paragraph("Contract Information", heading_style))
        metadata_data = [
            ["Document Name:", metadata.get("filename", "N/A")],
            ["Analysis Date:", metadata.get("analysis_date", datetime.now().strftime("%Y-%m-%d %H:%M"))],
            ["File Size:", metadata.get("file_size", "N/A")],
            ["Total Clauses:", str(len(clauses))]
        ]
        
        metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('BACKGROUND', (0, 0), (0, -1), colors.grey),
        ]))
        story.append(metadata_table)
        story.append(Spacer(1, 0.2*inch))
        
        # Executive Summary
        story.append(Paragraph("Executive Summary", heading_style))
        story.append(Paragraph(summaries.get("short_summary", "No summary available"), styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Risk Overview
        story.append(Paragraph("Risk Assessment Overview", heading_style))
        
        # Create risk chart
        risk_chart_path = _create_risk_chart(risk_summary)
        if risk_chart_path:
            story.append(Image(risk_chart_path, width=4*inch, height=3*inch))
            story.append(Spacer(1, 0.1*inch))
        
        # Risk summary text
        risk_level = risk_summary.get("contract_risk_level", "Unknown")
        risk_score = risk_summary.get("contract_risk_score", 0)
        story.append(Paragraph(
            f"<b>Overall Risk Level:</b> {risk_level}<br/>"
            f"<b>Risk Score:</b> {risk_score:.3f}<br/>"
            f"<b>High Risk Clauses:</b> {risk_summary.get('high_risk_count', 0)}<br/>"
            f"<b>Medium Risk Clauses:</b> {risk_summary.get('medium_risk_count', 0)}<br/>"
            f"<b>Low Risk Clauses:</b> {risk_summary.get('low_risk_count', 0)}",
            styles['Normal']
        ))
        story.append(Spacer(1, 0.2*inch))
        
        # Detailed Summary
        story.append(Paragraph("Detailed Analysis", heading_style))
        story.append(Paragraph(summaries.get("long_summary", "No detailed summary available"), styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Top Risky Clauses
        story.append(Paragraph("Top Risk Clauses", heading_style))
        top_risky = risk_summary.get("top_risky_clauses", [])[:5]
        
        if top_risky:
            risky_data = [["ID", "Type", "Risk Level", "Score", "Matched Terms"]]
            for clause_risk in top_risky:
                # Find the corresponding clause
                clause = next((c for c in clauses if c["id"] == clause_risk["clause_id"]), {})
                risky_data.append([
                    str(clause_risk["clause_id"]),
                    clause.get("type", "Unknown")[:15],
                    clause_risk["risk_level"],
                    f"{clause_risk['risk_score']:.3f}",
                    ", ".join(clause_risk["matched_terms"][:3])
                ])
            
            risky_table = Table(risky_data, colWidths=[0.5*inch, 1.2*inch, 1*inch, 0.8*inch, 2.5*inch])
            risky_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ]))
            story.append(risky_table)
        
        story.append(PageBreak())
        
        # All Clauses Table
        story.append(Paragraph("All Clauses", heading_style))
        
        # Create clauses table
        clause_data = [["ID", "Type", "Risk Level", "Entities", "Preview"]]
        for clause in clauses:
            # Find risk info for this clause
            clause_risk = next(
                (r for r in risk_summary.get("clause_risks", []) if r["clause_id"] == clause["id"]),
                {"risk_level": "Unknown"}
            )
            
            # Format entities
            entities_text = ", ".join([ent["text"] for ent in clause.get("entities", [])[:3]])
            if len(clause.get("entities", [])) > 3:
                entities_text += "..."
            
            clause_data.append([
                str(clause["id"]),
                clause["type"][:12],
                clause_risk["risk_level"],
                entities_text[:20],
                clause["text"][:50] + "..." if len(clause["text"]) > 50 else clause["text"]
            ])
        
        clause_table = Table(clause_data, colWidths=[0.4*inch, 1*inch, 0.8*inch, 1.3*inch, 2.5*inch])
        clause_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('GRID', (0, 0), (-1, -1), 1, colors.black),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ]))
        story.append(clause_table)
        
        # Footer
        story.append(Spacer(1, 0.3*inch))
        story.append(Paragraph(
            "<i>Generated by LawBrief AI - This analysis is for informational purposes only and does not constitute legal advice.</i>",
            styles['Normal']
        ))
        
        # Build PDF
        doc.build(story)
        
        # Cleanup temporary files
        if risk_chart_path and os.path.exists(risk_chart_path):
            os.remove(risk_chart_path)
            
        logging.info(f"PDF report generated successfully: {output_path}")
        
    except Exception as e:
        logging.error(f"Error generating PDF report: {e}")
        raise


def generate_docx_report(
    output_path: str,
    metadata: Dict,
    clauses: List[Dict],
    risk_summary: Dict, 
    summaries: Dict
) -> None:
    """
    Generate a professional DOCX report of contract analysis.
    
    Args:
        output_path: Path to save the DOCX report
        metadata: Contract metadata
        clauses: List of extracted clauses
        risk_summary: Risk assessment results
        summaries: Document summaries
    """
    try:
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('LawBrief AI - Contract Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contract Information
        doc.add_heading('Contract Information', level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ("Document Name:", metadata.get("filename", "N/A")),
            ("Analysis Date:", metadata.get("analysis_date", datetime.now().strftime("%Y-%m-%d %H:%M"))),
            ("File Size:", metadata.get("file_size", "N/A")),
            ("Total Clauses:", str(len(clauses)))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
            # Bold the labels
            info_table.cell(i, 0).paragraphs[0].runs.bold = True
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(summaries.get("short_summary", "No summary available"))
        
        # Risk Assessment Overview
        doc.add_heading('Risk Assessment Overview', level=1)
        
        risk_level = risk_summary.get("contract_risk_level", "Unknown")
        risk_score = risk_summary.get("contract_risk_score", 0)
        
        risk_para = doc.add_paragraph()
        risk_para.add_run("Overall Risk Level: ").bold = True
        risk_para.add_run(f"{risk_level}\n")
        risk_para.add_run("Risk Score: ").bold = True
        risk_para.add_run(f"{risk_score:.3f}\n")
        risk_para.add_run("High Risk Clauses: ").bold = True
        risk_para.add_run(f"{risk_summary.get('high_risk_count', 0)}\n")
        risk_para.add_run("Medium Risk Clauses: ").bold = True
        risk_para.add_run(f"{risk_summary.get('medium_risk_count', 0)}\n")
        risk_para.add_run("Low Risk Clauses: ").bold = True
        risk_para.add_run(f"{risk_summary.get('low_risk_count', 0)}")
        
        # Detailed Analysis
        doc.add_heading('Detailed Analysis', level=1)
        doc.add_paragraph(summaries.get("long_summary", "No detailed summary available"))
        
        # Top Risk Clauses
        doc.add_heading('Top Risk Clauses', level=1)
        top_risky = risk_summary.get("top_risky_clauses", [])[:5]
        
        if top_risky:
            risky_table = doc.add_table(rows=len(top_risky) + 1, cols=5)
            risky_table.style = 'Table Grid'
            
            # Headers
            headers = ["ID", "Type", "Risk Level", "Score", "Matched Terms"]
            for i, header in enumerate(headers):
                cell = risky_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs.bold = True
            
            # Data
            for row_idx, clause_risk in enumerate(top_risky, 1):
                clause = next((c for c in clauses if c["id"] == clause_risk["clause_id"]), {})
                risky_table.cell(row_idx, 0).text = str(clause_risk["clause_id"])
                risky_table.cell(row_idx, 1).text = clause.get("type", "Unknown")
                risky_table.cell(row_idx, 2).text = clause_risk["risk_level"]
                risky_table.cell(row_idx, 3).text = f"{clause_risk['risk_score']:.3f}"
                risky_table.cell(row_idx, 4).text = ", ".join(clause_risk["matched_terms"][:3])
        
        # Add page break
        doc.add_page_break()
        
        # All Clauses
        doc.add_heading('All Clauses', level=1)
        
        for clause in clauses:
            # Find risk info
            clause_risk = next(
                (r for r in risk_summary.get("clause_risks", []) if r["clause_id"] == clause["id"]),
                {"risk_level": "Unknown", "risk_score": 0}
            )
            
            # Clause heading
            clause_heading = doc.add_heading(f'Clause {clause["id"]}: {clause["type"]}', level=2)
            
            # Risk info
            risk_para = doc.add_paragraph()
            risk_para.add_run("Risk Level: ").bold = True
            risk_para.add_run(f"{clause_risk['risk_level']} ")
            risk_para.add_run("(Score: ").italic = True
            risk_para.add_run(f"{clause_risk.get('risk_score', 0):.3f}").italic = True
            risk_para.add_run(")").italic = True
            
            # Entities
            if clause.get("entities"):
                entities_para = doc.add_paragraph()
                entities_para.add_run("Entities: ").bold = True
                entity_texts = [f"{ent['text']} ({ent['label']})" for ent in clause["entities"]]
                entities_para.add_run(", ".join(entity_texts))
            
            # Clause text
            doc.add_paragraph(clause["text"])
            doc.add_paragraph("")  # Spacer
        
        # Footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            "Generated by LawBrief AI - This analysis is for informational purposes only and does not constitute legal advice."
        )
        footer_run.italic = True
        footer_run.font.size = Pt(10)
        
        # Save document
        doc.save(output_path)
        logging.info(f"DOCX report generated successfully: {output_path}")
        
    except Exception as e:
        logging.error(f"Error generating DOCX report: {e}")
        raise


def _create_risk_chart(risk_summary: Dict) -> str:
    """
    Create a risk distribution chart and return the file path.
    
    Args:
        risk_summary: Risk assessment summary
        
    Returns:
        Path to the generated chart image
    """
    try:
        risk_dist = risk_summary.get("risk_distribution", {"Low": 0, "Medium": 0, "High": 0})
        
        # Filter out zero values
        filtered_dist = {k: v for k, v in risk_dist.items() if v > 0}
        
        if not filtered_dist:
            return None
        
        # Create pie chart
        fig, ax = plt.subplots(figsize=(6, 4))
        
        labels = list(filtered_dist.keys())
        sizes = list(filtered_dist.values())
        colors = {'Low': '#90EE90', 'Medium': '#FFD700', 'High': '#FF6B6B'}
        chart_colors = [colors.get(label, '#CCCCCC') for label in labels]
        
        wedges, texts, autotexts = ax.pie(
            sizes, 
            labels=labels, 
            colors=chart_colors,
            autopct='%1.1f%%',
            startangle=90
        )
        
        # Style the chart
        ax.set_title('Risk Distribution', fontsize=14, fontweight='bold')
        
        # Make percentage text bold
        for autotext in autotexts:
            autotext.set_color('white')
            autotext.set_fontweight('bold')
        
        plt.tight_layout()
        
        # Save to temporary file
        temp_path = "temp_risk_chart.png"
        plt.savefig(temp_path, dpi=150, bbox_inches='tight')
        plt.close()
        
        return temp_path
        
    except Exception as e:
        logging.error(f"Error creating risk chart: {e}")
        return None


if __name__ == "__main__":
    # Demo functionality
    logging.basicConfig(level=logging.INFO)
    
    # Sample data for testing
    sample_metadata = {
        "filename": "sample_contract.pdf",
        "analysis_date": datetime.now().strftime("%Y-%m-%d %H:%M"),
        "file_size": "245 KB"
    }
    
    sample_clauses = [
        {
            "id": 1,
            "type": "Termination",
            "text": "This agreement may be terminated by either party with 30 days notice.",
            "entities": [{"text": "30 days", "label": "DATE", "start": 55, "end": 62}]
        },
        {
            "id": 2,
            "type": "Payment", 
            "text": "Company ABC shall pay $10,000 within 30 days of invoice.",
            "entities": [
                {"text": "Company ABC", "label": "ORG", "start": 0, "end": 11},
                {"text": "$10,000", "label": "MONEY", "start": 22, "end": 29}
            ]
        }
    ]
    
    sample_risk_summary = {
        "contract_risk_level": "Medium",
        "contract_risk_score": 0.45,
        "risk_distribution": {"Low": 1, "Medium": 1, "High": 0},
        "high_risk_count": 0,
        "medium_risk_count": 1, 
        "low_risk_count": 1,
        "clause_risks": [
            {"clause_id": 1, "risk_level": "Low", "risk_score": 0.2, "matched_terms": []},
            {"clause_id": 2, "risk_level": "Medium", "risk_score": 0.7, "matched_terms": ["payment"]}
        ],
        "top_risky_clauses": [
            {"clause_id": 2, "risk_level": "Medium", "risk_score": 0.7, "matched_terms": ["payment"]}
        ]
    }
    
    sample_summaries = {
        "short_summary": "This is a sample contract with basic termination and payment terms.",
        "long_summary": "This contract establishes a business relationship with standard termination and payment clauses. The overall risk level appears manageable with proper attention to payment terms."
    }
    
    print("Generating sample reports...")
    
    # Generate PDF
    try:
        generate_pdf_report(
            "sample_report.pdf",
            sample_metadata,
            sample_clauses,
            sample_risk_summary,
            sample_summaries
        )
        print("✓ PDF report generated")
    except Exception as e:
        print(f"✗ PDF generation failed: {e}")
    
    # Generate DOCX
    try:
        generate_docx_report(
            "sample_report.docx",
            sample_metadata, 
            sample_clauses,
            sample_risk_summary,
            sample_summaries
        )
        print("✓ DOCX report generated")
    except Exception as e:
        print(f"✗ DOCX generation failed: {e}")
